import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_code_block(doc, title, file_path):
    add_heading(doc, title, level=2)
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            # Adding code as a new paragraph with Courier font to simulate code blocks
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
    except Exception as e:
        doc.add_paragraph(f"Error loading file {file_path}: {str(e)}")

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('GUI Testing Report & Source Code', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta
    add_paragraph(doc, 'Project: AI Career Launchpad', bold=True)
    add_paragraph(doc, 'Date: June 12, 2026', bold=True)
    
    # 1. Overview
    add_heading(doc, '1. Overview', 1)
    doc.add_paragraph('Automated GUI tests have been successfully developed for the AI Career Launchpad application using Selenium WebDriver. The tests validate the critical paths of the frontend built with React, ensuring robust navigation, form interactions, authentication logic, and dynamic user interactions. This document includes both the report and the complete source code of the scripts developed.')
    
    # 2. Testing Scope & Scenarios Developed
    add_heading(doc, '2. Testing Scope & Scenarios Developed', 1)
    doc.add_paragraph('The automation suite addresses the primary functional requirements of the web application:')
    doc.add_paragraph('Authentication & Core (Auth-and-Core Suite): Validates the login process, verifying successful login leads to the dashboard and blocking invalid credentials. It also validates that core modules like the Resume Builder and Mock Interview modules save state effectively.', style='List Bullet')
    doc.add_paragraph('Navigation (Navigation Suite): Ensures that routing works properly, protected routes redirect unauthenticated users to /login, and the sidebar correctly links to all essential modules (Resume Builder, Skill Gap, Job Finder, Mock Interview, AI Chatbot).', style='List Bullet')
    doc.add_paragraph('Form Submissions (Forms Suite): End-to-end testing of profile data entry. Validates that updating user information (like Name and Bio) triggers the React state update and correctly saves via the backend API, displaying the proper success feedback.', style='List Bullet')
    doc.add_paragraph('User Interactions (Interactions Suite): Evaluates dynamic components, specifically the AI Chatbot rendering and the Job Finder modal. It confirms that the job details modal parses description text correctly and displays action buttons (Apply Now, Save).', style='List Bullet')

    # 3. Toolchain & Execution
    add_heading(doc, '3. Toolchain & Execution', 1)
    doc.add_paragraph('Selenium WebDriver (Node.js): Tests use selenium-webdriver with explicit wait conditions mapping to the frontend\'s data-testid properties and UI elements.', style='List Bullet')
    doc.add_paragraph('Selenium IDE: Created a recording scenario (login-dashboard.side) encapsulating a full test of the Login and Navigation flow directly from the browser extension.', style='List Bullet')
    doc.add_paragraph('Katalon Recorder: A supplementary Katalon WebDriver export (login-dashboard-webdriver.py) was generated to validate standard browser actions.', style='List Bullet')
    doc.add_paragraph('Selenium Grid Parallelization: The project incorporates a Docker-based Selenium Grid setup (docker-compose.grid.yml) allowing simultaneous execution of cross-browser tests across Chrome and Firefox nodes. This is executed using the npm run test:grid command.', style='List Bullet')

    # 4. Challenges Faced & Solutions Adopted
    add_heading(doc, '4. Challenges Faced & Solutions Adopted', 1)
    
    p = doc.add_paragraph(style='List Number')
    p.add_run('React State vs. Selenium Timing (Race Conditions):').bold = True
    doc.add_paragraph('Challenge: Selenium sendKeys operations sometimes triggered React synthetic events asynchronously, causing the "Save Changes" button to briefly evaluate as disabled.')
    doc.add_paragraph('Solution: Implemented explicit WebDriverWait conditions using XPath to target elements not just by text, but by their not(@disabled) state.')

    p = doc.add_paragraph(style='List Number')
    p.add_run('State Management with Local Storage:').bold = True
    doc.add_paragraph('Challenge: The standard driver.manage().deleteAllCookies() failed to effectively log the user out between test sessions because the Auth context persisted the JWT token inside localStorage.')
    doc.add_paragraph('Solution: Designed a clearSession() utility that navigates to the base URL and executes window.localStorage.clear(); window.sessionStorage.clear(); prior to clearing cookies.')

    p = doc.add_paragraph(style='List Number')
    p.add_run('Flaky UI Interactions (Modals & Sticky Buttons):').bold = True
    doc.add_paragraph('Challenge: The Job details modal used CSS text-transform: uppercase, failing initial text-matching assertions. The Profile Form\'s sticky "Save Changes" button frequently encountered ElementClickInterceptedException.')
    doc.add_paragraph('Solution: Adjusted assertions to match the computed DOM visibility state (e.g., matching "DESCRIPTION"). For the sticky button, we injected an arguments[0].scrollIntoView({block: "center"}) payload via executeScript before performing a standard .click().')

    # 5. Results Summary
    add_heading(doc, '5. Results Summary', 1)
    doc.add_paragraph('Passed Tests: Functional GUI checks successfully executed across Auth, Navigation, Forms, and Interactions.', style='List Bullet')
    doc.add_paragraph('Test Logs & Screenshots: Captured step-by-step UI states and stored them within the e2e/screenshots/report directory.', style='List Bullet')

    # 6. Selected Screenshots
    add_heading(doc, '6. Key Visual Evidence', 1)
    doc.add_paragraph('Below are selected screenshots captured during the automated test execution:')
    
    images = [
        ('02-dashboard-after-login.png', 'Successful Login and Dashboard Navigation'),
        ('05-resume-saved.png', 'Resume Builder Form Saving State'),
        ('09-sidebar-navigation.png', 'Sidebar Module Access Verification'),
        ('12-chatbot-reply.png', 'AI Chatbot Message Rendering'),
    ]
    
    base_img_path = os.path.join('e2e', 'screenshots', 'report')
    for img_name, caption in images:
        img_path = os.path.join(base_img_path, img_name)
        if os.path.exists(img_path):
            doc.add_paragraph()
            doc.add_picture(img_path, width=Inches(6.0))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            
    doc.add_page_break()

    # 7. Source Code Appendices
    add_heading(doc, '7. Source Code Appendices', 1)
    doc.add_paragraph('The following sections contain the raw source code for the automated scripts created to fulfill the GUI testing requirements.')

    # Requirement 1 & 2: Develop GUI tests and Create test scripts for different scenarios
    add_heading(doc, 'Requirement 1 & 2: Core Selenium WebDriver Test Scripts', 2)
    add_code_block(doc, 'Test Runner (e2e/run_all_tests.mjs)', os.path.join('e2e', 'run_all_tests.mjs'))
    add_code_block(doc, 'GUI Helpers (e2e/helpers/gui_helpers.mjs)', os.path.join('e2e', 'helpers', 'gui_helpers.mjs'))
    add_code_block(doc, 'Auth & Core Scenarios (e2e/tests/auth-and-core.test.mjs)', os.path.join('e2e', 'tests', 'auth-and-core.test.mjs'))
    add_code_block(doc, 'Navigation Scenarios (e2e/tests/navigation.test.mjs)', os.path.join('e2e', 'tests', 'navigation.test.mjs'))
    add_code_block(doc, 'Form Submission Scenarios (e2e/tests/forms.test.mjs)', os.path.join('e2e', 'tests', 'forms.test.mjs'))
    add_code_block(doc, 'User Interactions Scenarios (e2e/tests/interactions.test.mjs)', os.path.join('e2e', 'tests', 'interactions.test.mjs'))

    # Requirement 3: Explore Selenium Grid and Utilize it
    add_heading(doc, 'Requirement 3: Selenium Grid Parallel Testing', 2)
    add_code_block(doc, 'Docker Compose Grid Configuration (e2e/docker-compose.grid.yml)', os.path.join('e2e', 'docker-compose.grid.yml'))
    add_code_block(doc, 'Parallel Test Script (e2e/test_grid_parallel.mjs)', os.path.join('e2e', 'test_grid_parallel.mjs'))

    # Requirement 4: Implement at least one test using Selenium IDE and Katalon Recorder
    add_heading(doc, 'Requirement 4: Selenium IDE & Katalon Recorder', 2)
    add_code_block(doc, 'Selenium IDE Project File (e2e/selenium-ide/login-dashboard.side)', os.path.join('e2e', 'selenium-ide', 'login-dashboard.side'))
    
    katalon_path = os.path.join('e2e', 'katalon', 'login-dashboard-webdriver.py')
    if os.path.exists(katalon_path):
        add_code_block(doc, 'Katalon Recorder Python Export (e2e/katalon/login-dashboard-webdriver.py)', katalon_path)
    else:
        # Fallback to check if it's there or missing
        doc.add_paragraph('Katalon Recorder Python script path not found automatically.')

    # Save document
    doc.save('GUI_TESTING_REPORT.docx')
    print("Document successfully created at GUI_TESTING_REPORT.docx")

if __name__ == '__main__':
    create_report()
